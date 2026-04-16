from __future__ import annotations

import csv
import re
from pathlib import Path

from core.state import ContentSource


FIGMA_PATTERN = re.compile(r"https?://(?:www\.)?figma\.com/\S+", re.IGNORECASE)
SUPPORTED_DOCUMENT_SUFFIXES = {".pdf", ".csv", ".docx"}
QUOTED_PATH_PATTERN = re.compile(r"['\"]([^'\"]+\.(?:pdf|csv|docx))['\"]", re.IGNORECASE)
UNQUOTED_PATH_PATTERN = re.compile(r"(?<![\w/.-])([~./\w -]+\.(?:pdf|csv|docx))(?![\w/.-])", re.IGNORECASE)


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    return list(dict.fromkeys(item for item in items if item))


def _normalize_source(source: ContentSource) -> ContentSource:
    normalized = dict(source)
    normalized.setdefault("kind", "document")
    normalized.setdefault("uri", "")
    normalized.setdefault("display_name", normalized.get("uri", "source"))
    normalized.setdefault("format", "text")
    normalized.setdefault("extracted_text", "")
    normalized.setdefault("image_paths", [])
    normalized.setdefault("location_hints", [])
    normalized.setdefault("extraction_mode", "provided")
    normalized.setdefault("warnings", [])
    normalized["image_paths"] = list(normalized.get("image_paths", []))
    normalized["location_hints"] = list(normalized.get("location_hints", []))
    normalized["warnings"] = list(normalized.get("warnings", []))
    return normalized


def extract_document_paths(text: str) -> list[str]:
    quoted = QUOTED_PATH_PATTERN.findall(text)
    unquoted = [match.strip() for match in UNQUOTED_PATH_PATTERN.findall(text)]
    return _dedupe_preserve_order([*quoted, *unquoted])


def extract_figma_links(text: str) -> list[str]:
    return _dedupe_preserve_order(FIGMA_PATTERN.findall(text))


def resolve_document_path(raw_path: str, project_root: Path) -> Path:
    candidate = Path(raw_path.strip()).expanduser()
    if candidate.exists():
        return candidate.resolve()
    if candidate.is_absolute():
        return candidate
    return (project_root / candidate).resolve()


def has_document_source_hint(text: str, project_root: Path) -> bool:
    for raw_path in extract_document_paths(text):
        if resolve_document_path(raw_path, project_root).suffix.lower() in SUPPORTED_DOCUMENT_SUFFIXES:
            return True
    return False


def summarize_content_sources(sources: list[ContentSource]) -> list[str]:
    lines: list[str] = []
    for source in sources:
        display_name = source.get("display_name", source.get("uri", "source"))
        fmt = source.get("format", "text")
        mode = source.get("extraction_mode", "provided")
        warnings = source.get("warnings", [])
        warning_note = f" Warnings: {'; '.join(warnings)}" if warnings else ""
        lines.append(f"{display_name} [{fmt}] via {mode}.{warning_note}")
    return lines


def build_resolved_content_text(sources: list[ContentSource], user_text: str = "") -> str:
    sections: list[str] = []
    for source in sources:
        extracted_text = source.get("extracted_text", "").strip()
        if not extracted_text:
            continue
        display_name = source.get("display_name", source.get("uri", "source"))
        fmt = source.get("format", "text")
        sections.append(f"[Source: {display_name} | Format: {fmt}]\n{extracted_text}")

    request_text = user_text.strip()
    if request_text:
        sections.append(f"[Original Request]\n{request_text}")
    return "\n\n".join(sections).strip()


def resolve_content_sources(
    project_root: Path,
    user_text: str,
    provided_sources: list[ContentSource] | None = None,
) -> list[ContentSource]:
    sources: list[ContentSource] = []
    seen_uris: set[str] = set()

    for source in provided_sources or []:
        normalized = _normalize_source(source)
        key = normalized.get("uri", normalized.get("display_name", ""))
        if key in seen_uris:
            continue
        seen_uris.add(key)
        sources.append(normalized)

    for raw_path in extract_document_paths(user_text):
        resolved_path = resolve_document_path(raw_path, project_root)
        key = str(resolved_path)
        if key in seen_uris:
            continue
        seen_uris.add(key)
        sources.append(_extract_document_source(resolved_path))

    for link in extract_figma_links(user_text):
        if link in seen_uris:
            continue
        seen_uris.add(link)
        sources.append(
            _normalize_source(
                {
                    "kind": "figma",
                    "uri": link,
                    "display_name": link,
                    "format": "figma",
                    "extracted_text": "",
                    "image_paths": [],
                    "location_hints": [],
                    "extraction_mode": "unresolved",
                    "warnings": [
                        "Figma content review needs Codex/Figma pre-resolution or an exported artifact before the runtime can inspect the frame content."
                    ],
                }
            )
        )

    return sources


def _extract_document_source(path: Path) -> ContentSource:
    suffix = path.suffix.lower()
    base: ContentSource = {
        "kind": "document",
        "uri": str(path),
        "display_name": path.name,
        "format": suffix.lstrip(".") or "document",
        "extracted_text": "",
        "image_paths": [],
        "location_hints": [],
        "extraction_mode": "local-file",
        "warnings": [],
    }

    if not path.exists():
        base["warnings"] = [f"Local file was referenced but does not exist: {path}"]
        base["extraction_mode"] = "missing"
        return base

    if suffix == ".csv":
        return _extract_csv_source(path, base)
    if suffix == ".docx":
        return _extract_docx_source(path, base)
    if suffix == ".pdf":
        return _extract_pdf_source(path, base)

    base["warnings"] = [f"Unsupported document format for content ingestion: {suffix or 'unknown'}"]
    base["extraction_mode"] = "unsupported"
    return base


def _extract_csv_source(path: Path, base: ContentSource) -> ContentSource:
    rows: list[str] = []
    location_hints: list[str] = []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.reader(handle)
        parsed_rows = list(reader)

    if not parsed_rows:
        base["warnings"] = ["CSV file is empty."]
        base["extraction_mode"] = "empty"
        return base

    header = parsed_rows[0]
    rows.append(f"Header: {' | '.join(header)}")
    location_hints.append("row 1")
    for row_index, row in enumerate(parsed_rows[1:], start=2):
        rendered_cells = []
        for header_value, cell_value in zip(header, row):
            rendered_cells.append(f"{header_value}: {cell_value}")
        extra_cells = row[len(header) :]
        if extra_cells:
            rendered_cells.extend(f"extra_{index + 1}: {value}" for index, value in enumerate(extra_cells))
        rows.append(f"Row {row_index}: {' | '.join(rendered_cells)}")
        location_hints.append(f"row {row_index}")

    base["extracted_text"] = "\n".join(rows)
    base["location_hints"] = location_hints
    return base


def _extract_docx_source(path: Path, base: ContentSource) -> ContentSource:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        base["warnings"] = ["DOCX parsing requires python-docx to be installed."]
        base["extraction_mode"] = "dependency-missing"
        return base

    document = Document(str(path))
    lines: list[str] = []
    location_hints: list[str] = []

    paragraph_index = 1
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        lines.append(f"Paragraph {paragraph_index}: {text}")
        location_hints.append(f"paragraph {paragraph_index}")
        paragraph_index += 1

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            lines.append(f"Table {table_index}, Row {row_index}: {' | '.join(cells)}")
            location_hints.append(f"table {table_index} row {row_index}")

    if not lines:
        base["warnings"] = ["DOCX file did not contain extractable text paragraphs or tables."]
        base["extraction_mode"] = "empty"
        return base

    base["extracted_text"] = "\n".join(lines)
    base["location_hints"] = location_hints
    return base


def _extract_pdf_source(path: Path, base: ContentSource) -> ContentSource:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        base["warnings"] = ["PDF parsing requires pypdf to be installed."]
        base["extraction_mode"] = "dependency-missing"
        return base

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        base["warnings"] = [f"PDF could not be opened: {exc}"]
        base["extraction_mode"] = "open-failed"
        return base

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            base["warnings"] = ["PDF is encrypted and could not be decrypted for text extraction."]
            base["extraction_mode"] = "encrypted"
            return base

    lines: list[str] = []
    location_hints: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:
            text = ""
            base.setdefault("warnings", []).append(f"Failed to extract text from page {page_index}: {exc}")
        if not text:
            continue
        lines.append(f"Page {page_index}:\n{text}")
        location_hints.append(f"page {page_index}")

    if not lines:
        warnings = list(base.get("warnings", []))
        warnings.append("PDF did not expose extractable text. Provide a text-based PDF or export another artifact.")
        base["warnings"] = warnings
        base["extraction_mode"] = "no-text"
        return base

    base["extracted_text"] = "\n\n".join(lines)
    base["location_hints"] = location_hints
    return base
